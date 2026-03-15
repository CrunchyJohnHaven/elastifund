#!/usr/bin/env python3
"""
SystemIntel.docx Generator — Elastifund Signals Intelligence Brief
===================================================================
Reads from all available data sources (SQLite, JSON state, logs) and
produces a structured .docx intelligence document for LLM-guided research.

Data sources:
  - polymarket-bot/data/strategy_state_vps.json  (cycle stats, edge distribution)
  - polymarket-bot/data/paper_trades_vps.json     (paper trade records)
  - data/jj_trades.db                             (live trade database, if present)
  - data/quant.db                                 (research database)

Output:
  - SystemIntel.docx at project root

Usage:
    python scripts/generate_intel.py              # generate from local data
    python scripts/generate_intel.py --pull-vps   # SSH-pull latest from VPS first
"""

import json
import math
import os
import sqlite3
import statistics
import subprocess
import sys
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("ERROR: python-docx required. Install: pip3 install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(__file__).parent.parent
STRATEGY_STATE = ROOT / "polymarket-bot" / "data" / "strategy_state_vps.json"
PAPER_TRADES = ROOT / "polymarket-bot" / "data" / "paper_trades_vps.json"
JJ_TRADES_DB = ROOT / "data" / "jj_trades.db"
INTEL_SNAPSHOT = ROOT / "data" / "intel_snapshot.json"
QUANT_DB = ROOT / "data" / "quant.db"
OUTPUT = ROOT / "SystemIntel.docx"

# VPS connection details (for --pull-vps)
VPS_HOST = "ubuntu@52.208.155.0"
VPS_KEY = os.path.expanduser("~/Downloads/LightsailDefaultKey-eu-west-1.pem")
VPS_BOT_PATH = "/home/ubuntu/polymarket-trading-bot"


# ---------------------------------------------------------------------------
# Data loaders
# ---------------------------------------------------------------------------
def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    with open(path) as f:
        return json.load(f)


def load_db_rows(db_path: Path, query: str) -> list[dict]:
    if not db_path.exists():
        return []
    try:
        conn = sqlite3.connect(str(db_path))
        conn.row_factory = sqlite3.Row
        rows = conn.execute(query).fetchall()
        conn.close()
        return [dict(r) for r in rows]
    except Exception as e:
        print(f"  DB error ({db_path.name}): {e}")
        return []


def pull_vps_data():
    """SSH-pull latest data files from Dublin VPS."""
    print("Pulling data from VPS...")
    key_arg = f"-i {VPS_KEY}" if Path(VPS_KEY).exists() else ""
    pulls = [
        (f"{VPS_BOT_PATH}/data/jj_trades.db", str(JJ_TRADES_DB)),
        (f"{VPS_BOT_PATH}/data/intel_snapshot.json", str(INTEL_SNAPSHOT)),
        (f"{VPS_BOT_PATH}/jj_state.json", str(ROOT / "data" / "jj_state_vps.json")),
    ]
    for remote, local in pulls:
        Path(local).parent.mkdir(parents=True, exist_ok=True)
        cmd = f"scp -o StrictHostKeyChecking=no {key_arg} {VPS_HOST}:{remote} {local}"
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        if result.returncode == 0:
            print(f"  Pulled {Path(remote).name}")
        else:
            print(f"  Failed: {Path(remote).name} — {result.stderr.strip()}")


# ---------------------------------------------------------------------------
# Analysis functions
# ---------------------------------------------------------------------------
def analyze_edge_distribution(edges: list[float]) -> dict:
    if not edges:
        return {"count": 0}
    return {
        "count": len(edges),
        "mean": statistics.mean(edges),
        "median": statistics.median(edges),
        "stdev": statistics.stdev(edges) if len(edges) > 1 else 0,
        "p10": sorted(edges)[int(len(edges) * 0.1)],
        "p90": sorted(edges)[int(len(edges) * 0.9)],
        "min": min(edges),
        "max": max(edges),
        "pct_above_10": sum(1 for e in edges if e > 0.10) / len(edges) * 100,
        "pct_above_20": sum(1 for e in edges if e > 0.20) / len(edges) * 100,
        "pct_below_5": sum(1 for e in edges if e < 0.05) / len(edges) * 100,
    }


def analyze_paper_trades(data: dict) -> dict:
    portfolio = data.get("portfolio", {})
    open_pos = data.get("open_positions", [])
    closed = data.get("closed_trades", [])

    # Category classification (reuse jj_live logic)
    category_counts = Counter()
    category_edges = defaultdict(list)
    direction_counts = Counter()
    confidence_counts = Counter()
    edge_by_direction = defaultdict(list)

    all_trades = open_pos + closed
    for t in all_trades:
        q = (t.get("question") or "").lower()
        cat = _classify_question(q)
        category_counts[cat] += 1
        category_edges[cat].append(abs(t.get("edge", 0)))
        sig = t.get("signal", "UNKNOWN")
        direction_counts[sig] += 1
        edge_by_direction[sig].append(abs(t.get("edge", 0)))
        confidence_counts[t.get("confidence", "unknown")] += 1

    # Win/loss analysis on closed trades
    wins = sum(1 for t in closed if (t.get("pnl") or 0) > 0)
    losses = sum(1 for t in closed if (t.get("pnl") or 0) < 0)
    total_pnl = sum(t.get("pnl", 0) or 0 for t in closed)

    return {
        "portfolio": portfolio,
        "open_count": len(open_pos),
        "closed_count": len(closed),
        "wins": wins,
        "losses": losses,
        "win_rate": wins / max(1, wins + losses) * 100,
        "total_pnl": total_pnl,
        "category_counts": dict(category_counts.most_common()),
        "category_avg_edge": {
            cat: statistics.mean(edges) for cat, edges in category_edges.items()
        },
        "direction_counts": dict(direction_counts),
        "confidence_counts": dict(confidence_counts),
        "edge_by_direction": {
            d: statistics.mean(edges) for d, edges in edge_by_direction.items()
        },
    }


def analyze_live_trades(db_path: Path) -> dict:
    trades = load_db_rows(db_path, "SELECT * FROM trades ORDER BY timestamp DESC")
    cycles = load_db_rows(db_path, "SELECT * FROM cycles ORDER BY timestamp DESC LIMIT 50")
    daily = load_db_rows(db_path, "SELECT * FROM daily_reports ORDER BY date DESC LIMIT 30")

    if not trades and not cycles:
        return {"has_data": False}

    # Trade analysis
    resolved = [t for t in trades if t.get("outcome")]
    wins = [t for t in resolved if t.get("outcome") == "won"]
    losses = [t for t in resolved if t.get("outcome") == "lost"]

    # Category breakdown
    cat_stats = defaultdict(lambda: {"count": 0, "wins": 0, "losses": 0, "pnl": 0.0, "edges": []})
    for t in trades:
        cat = t.get("category", "unknown") or "unknown"
        cat_stats[cat]["count"] += 1
        cat_stats[cat]["edges"].append(abs(t.get("edge", 0) or 0))
        if t.get("outcome") == "won":
            cat_stats[cat]["wins"] += 1
        elif t.get("outcome") == "lost":
            cat_stats[cat]["losses"] += 1
        cat_stats[cat]["pnl"] += t.get("pnl", 0) or 0

    # Direction analysis
    yes_trades = [t for t in trades if t.get("direction") == "buy_yes"]
    no_trades = [t for t in trades if t.get("direction") == "buy_no"]

    # Calibration: compare estimated prob to actual outcome
    calibration_buckets = defaultdict(lambda: {"count": 0, "wins": 0})
    for t in resolved:
        prob = t.get("calibrated_prob") or t.get("raw_prob")
        if prob is None:
            continue
        bucket = round(prob * 10) / 10  # 0.0, 0.1, ..., 1.0
        calibration_buckets[bucket]["count"] += 1
        if t.get("outcome") == "won":
            calibration_buckets[bucket]["wins"] += 1

    # Cycle efficiency
    cycle_stats = {}
    if cycles:
        cycle_stats = {
            "recent_count": len(cycles),
            "avg_markets_scanned": statistics.mean(
                [c.get("markets_scanned", 0) or 0 for c in cycles]
            ),
            "avg_signals": statistics.mean(
                [c.get("signals_found", 0) or 0 for c in cycles]
            ),
            "avg_trades": statistics.mean(
                [c.get("trades_placed", 0) or 0 for c in cycles]
            ),
            "signal_to_trade_rate": (
                sum(c.get("trades_placed", 0) or 0 for c in cycles)
                / max(1, sum(c.get("signals_found", 0) or 0 for c in cycles))
                * 100
            ),
        }

    return {
        "has_data": True,
        "total_trades": len(trades),
        "resolved": len(resolved),
        "wins": len(wins),
        "losses": len(losses),
        "win_rate": len(wins) / max(1, len(resolved)) * 100,
        "total_pnl": sum(t.get("pnl", 0) or 0 for t in resolved),
        "yes_count": len(yes_trades),
        "no_count": len(no_trades),
        "category_stats": {k: dict(v) for k, v in cat_stats.items()},
        "calibration": dict(calibration_buckets),
        "cycle_stats": cycle_stats,
        "daily_reports": daily[:7],  # Last 7 days
    }


def _classify_question(question: str) -> str:
    """Lightweight category classifier (mirrors jj_live.py logic)."""
    q = question.lower()
    categories = {
        "politics": ["election", "president", "congress", "senate", "trump", "biden",
                      "democrat", "republican", "vote", "governor"],
        "weather": ["temperature", "rain", "weather", "hurricane", "storm", "degrees",
                     "fahrenheit", "celsius", "high of", "low of"],
        "crypto": ["bitcoin", "btc", "ethereum", "eth", "crypto", "solana", "token",
                    "blockchain", "memecoin"],
        "sports": ["nba", "nfl", "mlb", "nhl", "soccer", "football", "basketball",
                    "championship", "playoff", "world cup", "ufc", "qualify",
                    "premier league", "champions league", "fifa"],
        "geopolitical": ["war", "invasion", "nato", "china", "russia", "taiwan",
                          "ceasefire", "nuclear", "military", "invades"],
        "economic": ["inflation", "cpi", "gdp", "unemployment", "fed", "fomc",
                      "interest rate", "recession", "treasury"],
    }
    scores = {}
    for cat, keywords in categories.items():
        scores[cat] = sum(1 for kw in keywords if kw in q)
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "other"


def generate_research_directives(
    edge_stats: dict,
    paper_analysis: dict,
    live_analysis: dict,
    strategy_state: dict,
) -> list[str]:
    """Auto-generate research priorities from the data."""
    directives = []

    # 1. Edge concentration analysis
    if edge_stats.get("pct_above_20", 0) > 30:
        directives.append(
            f"INVESTIGATE: {edge_stats['pct_above_20']:.0f}% of signals show >20% edge. "
            "This is suspiciously high — likely indicates calibration error or "
            "systematic mispricing in a specific category. Decompose by category."
        )

    if edge_stats.get("pct_below_5", 0) > 50:
        directives.append(
            "EFFICIENCY: >50% of signals have <5% edge (below noise threshold). "
            "Investigate whether scanner filters should be tightened to reduce "
            "API calls and focus LLM analysis on higher-probability opportunities."
        )

    # 2. Category gaps
    cat_counts = paper_analysis.get("category_counts", {})
    cat_edges = paper_analysis.get("category_avg_edge", {})
    if cat_counts:
        best_cat = max(cat_edges, key=cat_edges.get) if cat_edges else None
        if best_cat and cat_edges[best_cat] > 0.15:
            directives.append(
                f"DOUBLE DOWN: '{best_cat}' category shows highest avg edge "
                f"({cat_edges[best_cat]:.1%}). Research deeper edge sources: "
                f"is it structural (data advantage) or behavioral (market bias)?"
            )

    # 3. Direction asymmetry
    dir_edges = paper_analysis.get("edge_by_direction", {})
    if "SELL" in dir_edges and "BUY" in dir_edges:
        sell_edge = dir_edges["SELL"]
        buy_edge = dir_edges["BUY"]
        if sell_edge > buy_edge * 1.5:
            directives.append(
                f"ASYMMETRY: NO/SELL signals show {sell_edge:.1%} avg edge vs "
                f"BUY at {buy_edge:.1%}. Consistent with research (NO outperforms "
                "YES at 69/99 price levels). Bias allocation toward NO positions."
            )

    # 4. Signal volume trends
    total_signals = strategy_state.get("total_signals", 0)
    cycles = strategy_state.get("cycles_completed", 0)
    if cycles > 0:
        signals_per_cycle = total_signals / cycles
        directives.append(
            f"THROUGHPUT: {signals_per_cycle:.1f} signals/cycle across {cycles} cycles. "
            f"Total signal inventory: {total_signals:,}. "
            "Track whether signal quality improves as calibration data accumulates."
        )

    # 5. Confidence distribution
    conf = strategy_state.get("signals_by_confidence", {})
    if conf:
        total = sum(conf.values())
        high_pct = conf.get("high", 0) / max(1, total) * 100
        directives.append(
            f"CONFIDENCE MIX: {high_pct:.0f}% high-confidence signals "
            f"(high={conf.get('high',0)}, med={conf.get('medium',0)}, "
            f"low={conf.get('low',0)}). "
            "Validate that high-confidence signals actually outperform."
        )

    # 6. Live performance gaps
    if live_analysis.get("has_data"):
        wr = live_analysis.get("win_rate", 0)
        if wr > 0 and wr < 55:
            directives.append(
                f"WARNING: Live win rate at {wr:.1f}% — below profitability "
                "threshold. Audit recent losses for systematic patterns."
            )
        elif wr >= 60:
            directives.append(
                f"POSITIVE: Live win rate at {wr:.1f}%. Validate this isn't "
                "survivorship bias from small sample. Need 50+ resolved trades "
                "for statistical significance."
            )

    # 7. Calibration health
    if live_analysis.get("calibration"):
        cal = live_analysis["calibration"]
        overconfident = sum(
            1 for bucket, stats in cal.items()
            if bucket > 0.5 and stats["count"] > 0
            and stats["wins"] / stats["count"] < bucket - 0.1
        )
        if overconfident > 0:
            directives.append(
                f"CALIBRATION DRIFT: {overconfident} probability buckets show "
                "overconfidence (predicted > actual). Consider tightening "
                "Platt scaling parameters or increasing calibration penalty."
            )

    if not directives:
        directives.append(
            "BASELINE: Insufficient resolved trade data for auto-generated "
            "directives. Priority: accumulate 50+ resolved trades across "
            "3+ categories before drawing conclusions."
        )

    return directives


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------
def build_document(
    strategy_state: dict,
    paper_analysis: dict,
    live_analysis: dict,
    edge_stats: dict,
    directives: list[str],
    intel_snapshot: dict | None = None,
) -> Document:
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Title
    title = doc.add_heading("ELASTIFUND SIGNALS INTELLIGENCE BRIEF", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.now(timezone.utc)
    meta.add_run(f"Generated: {now.strftime('%Y-%m-%d %H:%M UTC')}\n").bold = False
    meta.add_run(f"Cycles analyzed: {strategy_state.get('cycles_completed', 0)}\n")
    meta.add_run(f"Total signals processed: {strategy_state.get('total_signals', 0):,}\n")
    meta.add_run("Classification: INTERNAL — Feed to LLM research sessions")

    doc.add_page_break()

    # =====================================================================
    # SECTION 1: EXECUTIVE DASHBOARD
    # =====================================================================
    doc.add_heading("1. EXECUTIVE DASHBOARD", level=1)

    portfolio = paper_analysis.get("portfolio", {})

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    hdr[2].text = "Metric"
    hdr[3].text = "Value"

    dashboard_data = [
        ("Starting Capital", f"${portfolio.get('starting_capital', 0):.0f}"),
        ("Cash Available", f"${portfolio.get('cash', 0):.0f}"),
        ("Total Trades", str(portfolio.get('total_trades', 0))),
        ("Win Rate", f"{portfolio.get('win_rate', 0):.1f}%"),
        ("Open Positions", str(paper_analysis.get('open_count', 0))),
        ("Realized PnL", f"${portfolio.get('realized_pnl', 0):.2f}"),
        ("Signals Processed", f"{strategy_state.get('total_signals', 0):,}"),
        ("Cycles Run", str(strategy_state.get('cycles_completed', 0))),
    ]

    for i in range(0, len(dashboard_data), 2):
        row = table.add_row().cells
        row[0].text = dashboard_data[i][0]
        row[1].text = dashboard_data[i][1]
        if i + 1 < len(dashboard_data):
            row[2].text = dashboard_data[i + 1][0]
            row[3].text = dashboard_data[i + 1][1]

    # Live trading status
    if live_analysis.get("has_data"):
        doc.add_heading("Live Trading Performance", level=2)
        live_p = doc.add_paragraph()
        live_p.add_run(f"Total live trades: {live_analysis['total_trades']}\n")
        live_p.add_run(f"Resolved: {live_analysis['resolved']} ")
        live_p.add_run(f"(W: {live_analysis['wins']} / L: {live_analysis['losses']})\n")
        live_p.add_run(f"Win rate: {live_analysis['win_rate']:.1f}%\n")
        live_p.add_run(f"Cumulative PnL: ${live_analysis['total_pnl']:.2f}\n")
        live_p.add_run(f"YES trades: {live_analysis['yes_count']} | ")
        live_p.add_run(f"NO trades: {live_analysis['no_count']}")

    # =====================================================================
    # SECTION 2: SIGNAL FLOW ANALYSIS
    # =====================================================================
    doc.add_heading("2. SIGNAL FLOW ANALYSIS", level=1)

    doc.add_heading("Edge Distribution (All Observed Signals)", level=2)

    if edge_stats.get("count", 0) > 0:
        table2 = doc.add_table(rows=1, cols=2)
        table2.style = "Light Shading Accent 1"
        table2.rows[0].cells[0].text = "Statistic"
        table2.rows[0].cells[1].text = "Value"

        edge_rows = [
            ("Sample Size", f"{edge_stats['count']}"),
            ("Mean Edge", f"{edge_stats['mean']:.1%}"),
            ("Median Edge", f"{edge_stats['median']:.1%}"),
            ("Std Dev", f"{edge_stats['stdev']:.1%}"),
            ("10th Percentile", f"{edge_stats['p10']:.1%}"),
            ("90th Percentile", f"{edge_stats['p90']:.1%}"),
            ("Min / Max", f"{edge_stats['min']:.1%} / {edge_stats['max']:.1%}"),
            ("Signals >10% edge", f"{edge_stats['pct_above_10']:.1f}%"),
            ("Signals >20% edge", f"{edge_stats['pct_above_20']:.1f}%"),
            ("Signals <5% edge", f"{edge_stats['pct_below_5']:.1f}%"),
        ]
        for label, val in edge_rows:
            row = table2.add_row().cells
            row[0].text = label
            row[1].text = val
    else:
        doc.add_paragraph("No edge distribution data available.")

    # Signal confidence breakdown
    conf = strategy_state.get("signals_by_confidence", {})
    if conf:
        doc.add_heading("Confidence Distribution", level=2)
        total_signals = sum(conf.values())
        table3 = doc.add_table(rows=1, cols=3)
        table3.style = "Light Shading Accent 1"
        table3.rows[0].cells[0].text = "Confidence"
        table3.rows[0].cells[1].text = "Count"
        table3.rows[0].cells[2].text = "Percentage"
        for level in ["high", "medium", "low"]:
            count = conf.get(level, 0)
            row = table3.add_row().cells
            row[0].text = level.upper()
            row[1].text = str(count)
            row[2].text = f"{count / max(1, total_signals) * 100:.1f}%"

    # =====================================================================
    # SECTION 3: CATEGORY INTELLIGENCE
    # =====================================================================
    doc.add_heading("3. CATEGORY INTELLIGENCE", level=1)

    cat_counts = paper_analysis.get("category_counts", {})
    cat_edges = paper_analysis.get("category_avg_edge", {})

    if cat_counts:
        table4 = doc.add_table(rows=1, cols=4)
        table4.style = "Light Shading Accent 1"
        hdr = table4.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Trades"
        hdr[2].text = "Avg Edge"
        hdr[3].text = "Assessment"

        for cat in sorted(cat_counts, key=cat_counts.get, reverse=True):
            avg_edge = cat_edges.get(cat, 0)
            # Assessment logic
            if avg_edge > 0.20:
                assessment = "HIGH PRIORITY — Large edge, validate not noise"
            elif avg_edge > 0.10:
                assessment = "PROMISING — Monitor for consistency"
            elif avg_edge > 0.05:
                assessment = "MARGINAL — Edge near noise floor"
            else:
                assessment = "LOW VALUE — Below minimum threshold"

            row = table4.add_row().cells
            row[0].text = cat.upper()
            row[1].text = str(cat_counts[cat])
            row[2].text = f"{avg_edge:.1%}"
            row[3].text = assessment

    # Direction analysis
    doc.add_heading("Direction Asymmetry", level=2)
    dir_counts = paper_analysis.get("direction_counts", {})
    dir_edges = paper_analysis.get("edge_by_direction", {})
    if dir_counts:
        p = doc.add_paragraph()
        for direction, count in dir_counts.items():
            avg = dir_edges.get(direction, 0)
            p.add_run(f"{direction}: {count} trades, avg edge {avg:.1%}\n")

    # Live category performance
    if live_analysis.get("has_data") and live_analysis.get("category_stats"):
        doc.add_heading("Live Category Performance", level=2)
        cat_stats = live_analysis["category_stats"]
        table5 = doc.add_table(rows=1, cols=5)
        table5.style = "Light Shading Accent 1"
        hdr = table5.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Trades"
        hdr[2].text = "Wins"
        hdr[3].text = "Win Rate"
        hdr[4].text = "PnL"

        for cat, stats in sorted(cat_stats.items(), key=lambda x: x[1]["pnl"], reverse=True):
            wr = stats["wins"] / max(1, stats["wins"] + stats["losses"]) * 100
            row = table5.add_row().cells
            row[0].text = cat.upper()
            row[1].text = str(stats["count"])
            row[2].text = str(stats["wins"])
            row[3].text = f"{wr:.0f}%"
            row[4].text = f"${stats['pnl']:.2f}"

    # =====================================================================
    # SECTION 4: CALIBRATION HEALTH
    # =====================================================================
    doc.add_heading("4. CALIBRATION HEALTH", level=1)

    p = doc.add_paragraph()
    p.add_run("Platt Scaling Parameters: ").bold = True
    p.add_run(f"A={0.55}, B={-0.40}\n")
    p.add_run("Training: 70% of 532 resolved markets\n")
    p.add_run("Test Brier: 0.286 (raw) → 0.245 (calibrated)\n")
    p.add_run("Mapping: 90% → 71%, 80% → 60%, 70% → 53%\n\n")

    if live_analysis.get("calibration"):
        doc.add_heading("Predicted vs Actual (Live Data)", level=2)
        cal = live_analysis["calibration"]
        table6 = doc.add_table(rows=1, cols=4)
        table6.style = "Light Shading Accent 1"
        hdr = table6.rows[0].cells
        hdr[0].text = "Predicted Prob"
        hdr[1].text = "Trades"
        hdr[2].text = "Actual Win %"
        hdr[3].text = "Drift"

        for bucket in sorted(cal.keys()):
            stats = cal[bucket]
            if stats["count"] == 0:
                continue
            actual = stats["wins"] / stats["count"] * 100
            drift = actual - bucket * 100
            row = table6.add_row().cells
            row[0].text = f"{bucket:.0%}"
            row[1].text = str(stats["count"])
            row[2].text = f"{actual:.0f}%"
            row[3].text = f"{drift:+.0f}pp"
    else:
        doc.add_paragraph(
            "Insufficient resolved trade data for calibration analysis. "
            "Need 50+ resolved trades with probability estimates."
        )

    # =====================================================================
    # SECTION 5: POSITION INVENTORY
    # =====================================================================
    doc.add_heading("5. OPEN POSITION INVENTORY", level=1)

    paper_data = load_json(PAPER_TRADES)
    open_pos = paper_data.get("open_positions", [])

    if open_pos:
        table7 = doc.add_table(rows=1, cols=5)
        table7.style = "Light Shading Accent 1"
        hdr = table7.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Question (truncated)"
        hdr[2].text = "Direction"
        hdr[3].text = "Edge"
        hdr[4].text = "Confidence"

        for pos in sorted(open_pos, key=lambda x: abs(x.get("edge", 0)), reverse=True):
            row = table7.add_row().cells
            row[0].text = pos.get("trade_id", "?")
            q = pos.get("question", "")
            row[1].text = q[:55] + "..." if len(q) > 55 else q
            row[2].text = pos.get("signal", "?")
            row[3].text = f"{abs(pos.get('edge', 0)):.1%}"
            row[4].text = str(pos.get("confidence", "?"))

        doc.add_paragraph(
            f"\nTotal exposure: ${sum(p.get('size_usdc', 0) for p in open_pos):.2f}"
        )
    else:
        doc.add_paragraph("No open positions.")

    # =====================================================================
    # SECTION 6: ANOMALY DETECTION
    # =====================================================================
    doc.add_heading("6. ANOMALY DETECTION", level=1)

    anomalies = []

    # Check for edge clustering (many signals at same edge level = possible bug)
    edges = strategy_state.get("edge_distribution", [])
    if edges:
        edge_counter = Counter(round(e, 4) for e in edges)
        most_common_edge, count = edge_counter.most_common(1)[0]
        if count > len(edges) * 0.05:
            anomalies.append(
                f"EDGE CLUSTERING: {count} signals ({count/len(edges)*100:.0f}%) "
                f"at edge={most_common_edge:.4f}. Possible systematic artifact."
            )

    # Check for category dominance
    if cat_counts:
        total_trades = sum(cat_counts.values())
        for cat, count in cat_counts.items():
            if count / total_trades > 0.4:
                anomalies.append(
                    f"CATEGORY DOMINANCE: '{cat}' represents {count/total_trades*100:.0f}% "
                    f"of all trades. Portfolio is underdiversified."
                )

    # Check for extreme edges (possible data errors)
    extreme_edges = [e for e in edges if e > 0.50]
    if extreme_edges:
        anomalies.append(
            f"EXTREME EDGES: {len(extreme_edges)} signals show >50% edge. "
            "These are almost certainly calibration artifacts or data errors, "
            "not real mispricing."
        )

    if anomalies:
        for a in anomalies:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(a)
    else:
        doc.add_paragraph("No anomalies detected in current data.")

    # =====================================================================
    # SECTION 7: REAL-TIME CYCLE TRENDS (from intel snapshot)
    # =====================================================================
    doc.add_heading("7. REAL-TIME CYCLE TRENDS", level=1)

    if intel_snapshot and intel_snapshot.get("cycle_history"):
        history = intel_snapshot["cycle_history"]
        recent = history[-20:]  # Last 20 cycles

        doc.add_heading("Recent Cycle Efficiency", level=2)
        table_cyc = doc.add_table(rows=1, cols=6)
        table_cyc.style = "Light Shading Accent 1"
        hdr = table_cyc.rows[0].cells
        hdr[0].text = "Cycle"
        hdr[1].text = "Scanned"
        hdr[2].text = "Actionable"
        hdr[3].text = "Signals"
        hdr[4].text = "Trades"
        hdr[5].text = "Bankroll"

        for cyc in recent[-10:]:  # Show last 10 in table
            row = table_cyc.add_row().cells
            row[0].text = str(cyc.get("cycle", "?"))
            row[1].text = str(cyc.get("markets_scanned", 0))
            row[2].text = str(cyc.get("markets_actionable", 0))
            row[3].text = str(cyc.get("signals_generated", 0))
            row[4].text = str(cyc.get("trades_placed", 0))
            row[5].text = f"${cyc.get('bankroll', 0):.0f}"

        # Trend analysis
        if len(recent) >= 5:
            avg_signals = statistics.mean([c.get("signals_generated", 0) for c in recent])
            avg_trades = statistics.mean([c.get("trades_placed", 0) for c in recent])
            conversion = avg_trades / max(0.01, avg_signals) * 100

            first_half = recent[:len(recent)//2]
            second_half = recent[len(recent)//2:]
            sig_trend = (
                statistics.mean([c.get("signals_generated", 0) for c in second_half])
                - statistics.mean([c.get("signals_generated", 0) for c in first_half])
            )

            p = doc.add_paragraph()
            p.add_run("Trend Summary: ").bold = True
            p.add_run(
                f"Avg {avg_signals:.1f} signals/cycle, "
                f"{avg_trades:.1f} trades/cycle, "
                f"{conversion:.0f}% signal-to-trade conversion. "
                f"Signal volume {'increasing' if sig_trend > 0 else 'decreasing'} "
                f"({sig_trend:+.1f}/cycle vs prior window)."
            )

        # Recent signal patterns
        recent_sigs = intel_snapshot.get("recent_signals", [])
        if recent_sigs:
            doc.add_heading("Recent Signal Patterns", level=2)

            # Category breakdown of recent signals
            cat_counts = Counter(s.get("category", "unknown") for s in recent_sigs)
            dir_counts = Counter(s.get("direction", "hold") for s in recent_sigs)
            recent_edges = [s.get("edge", 0) for s in recent_sigs if s.get("edge", 0) > 0]

            p = doc.add_paragraph()
            p.add_run(f"Last {len(recent_sigs)} signals:\n")
            p.add_run("Categories: ")
            p.add_run(", ".join(f"{cat}: {n}" for cat, n in cat_counts.most_common(5)))
            p.add_run("\nDirections: ")
            p.add_run(", ".join(f"{d}: {n}" for d, n in dir_counts.most_common()))
            if recent_edges:
                p.add_run(f"\nAvg edge: {statistics.mean(recent_edges):.1%}")
                p.add_run(f", Median: {statistics.median(recent_edges):.1%}")

            # Top 5 highest-edge recent signals
            top_sigs = sorted(recent_sigs, key=lambda s: abs(s.get("edge", 0)), reverse=True)[:5]
            if top_sigs:
                doc.add_heading("Highest-Edge Recent Signals", level=3)
                for sig in top_sigs:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(
                        f"[{sig.get('direction', '?')}] edge={abs(sig.get('edge', 0)):.1%} "
                        f"cat={sig.get('category', '?')} — {sig.get('question', '?')}"
                    )

        # Category frequency map
        cat_freq = intel_snapshot.get("category_frequency", {})
        if cat_freq:
            doc.add_heading("Lifetime Category Distribution", level=2)
            total = sum(cat_freq.values())
            table_cf = doc.add_table(rows=1, cols=3)
            table_cf.style = "Light Shading Accent 1"
            table_cf.rows[0].cells[0].text = "Category"
            table_cf.rows[0].cells[1].text = "Signals"
            table_cf.rows[0].cells[2].text = "Share"
            for cat, count in sorted(cat_freq.items(), key=lambda x: x[1], reverse=True):
                row = table_cf.add_row().cells
                row[0].text = cat.upper()
                row[1].text = str(count)
                row[2].text = f"{count / max(1, total) * 100:.1f}%"
    else:
        doc.add_paragraph(
            "No real-time intel snapshot available. Deploy jj_live.py with the "
            "intel snapshot writer, then pull with: python scripts/generate_intel.py --pull-vps"
        )

    # =====================================================================
    # SECTION 8: RESEARCH DIRECTIVES (AUTO-GENERATED)
    # =====================================================================
    doc.add_heading("8. RESEARCH DIRECTIVES", level=1)

    doc.add_paragraph(
        "The following research priorities are auto-generated from observed data patterns. "
        "Feed this section to Deep Research prompts to guide the next investigation cycle."
    )

    for i, directive in enumerate(directives, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"D{i}. ")
        run.bold = True
        p.add_run(directive)

    # =====================================================================
    # SECTION 9: SYSTEM PARAMETERS (CURRENT)
    # =====================================================================
    doc.add_heading("9. SYSTEM PARAMETERS", level=1)

    params = {
        "Max Position USD": "$15",
        "Max Daily Loss USD": "$25",
        "Max Exposure": "90%",
        "Kelly Fraction": "0.50 (half-Kelly)",
        "Scan Interval": "180s (3 min)",
        "Max Open Positions": "30",
        "Min Edge Threshold": "5%",
        "Max Resolution Hours": "48h",
        "YES Threshold": "15%",
        "NO Threshold": "5%",
        "Platt A / B": "0.55 / -0.40",
        "Maker Fees": "0% (post-only orders)",
        "Categories Skipped": "crypto, sports, fed_rates, financial_speculation",
    }

    table8 = doc.add_table(rows=1, cols=2)
    table8.style = "Light Shading Accent 1"
    table8.rows[0].cells[0].text = "Parameter"
    table8.rows[0].cells[1].text = "Value"
    for param, val in params.items():
        row = table8.add_row().cells
        row[0].text = param
        row[1].text = val

    # =====================================================================
    # SECTION 10: DATA FRESHNESS
    # =====================================================================
    doc.add_heading("10. DATA FRESHNESS & SOURCES", level=1)

    sources = [
        ("Strategy State (VPS)", str(STRATEGY_STATE), STRATEGY_STATE.exists()),
        ("Paper Trades (VPS)", str(PAPER_TRADES), PAPER_TRADES.exists()),
        ("Intel Snapshot", str(INTEL_SNAPSHOT), INTEL_SNAPSHOT.exists()),
        ("Live Trades DB", str(JJ_TRADES_DB), JJ_TRADES_DB.exists()),
        ("Research DB (quant)", str(QUANT_DB), QUANT_DB.exists()),
    ]

    table9 = doc.add_table(rows=1, cols=3)
    table9.style = "Light Shading Accent 1"
    table9.rows[0].cells[0].text = "Source"
    table9.rows[0].cells[1].text = "Path"
    table9.rows[0].cells[2].text = "Available"
    for name, path, avail in sources:
        row = table9.add_row().cells
        row[0].text = name
        row[1].text = Path(path).name
        row[2].text = "YES" if avail else "NO"

    doc.add_paragraph(
        f"\nDocument generated: {now.strftime('%Y-%m-%d %H:%M:%S UTC')}\n"
        "Regenerate with: python scripts/generate_intel.py"
    )

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("=" * 60)
    print("ELASTIFUND SIGNALS INTELLIGENCE GENERATOR")
    print("=" * 60)

    # Optional VPS pull
    if "--pull-vps" in sys.argv:
        pull_vps_data()

    # Load data
    print("\nLoading data sources...")
    strategy_state = load_json(STRATEGY_STATE)
    print(f"  Strategy state: {strategy_state.get('cycles_completed', 0)} cycles")

    intel_snapshot = load_json(INTEL_SNAPSHOT)
    if intel_snapshot:
        print(f"  Intel snapshot: updated {intel_snapshot.get('last_updated', 'unknown')}")
        print(f"    Recent signals: {len(intel_snapshot.get('recent_signals', []))}")
        print(f"    Cycle history: {len(intel_snapshot.get('cycle_history', []))}")
        # Merge snapshot data into strategy_state for richer analysis
        if not strategy_state.get("cycles_completed"):
            strategy_state["cycles_completed"] = intel_snapshot.get("total_cycles", 0)
        if not strategy_state.get("total_signals"):
            strategy_state["total_signals"] = intel_snapshot.get("total_signals", 0)
        if not strategy_state.get("signals_by_confidence") and intel_snapshot.get("signals_by_confidence"):
            strategy_state["signals_by_confidence"] = intel_snapshot["signals_by_confidence"]

    paper_data = load_json(PAPER_TRADES)
    paper_analysis = analyze_paper_trades(paper_data)
    print(f"  Paper trades: {paper_analysis.get('open_count', 0)} open, "
          f"{paper_analysis.get('closed_count', 0)} closed")

    live_analysis = analyze_live_trades(JJ_TRADES_DB)
    print(f"  Live trades DB: {'found' if live_analysis.get('has_data') else 'no data'}")

    edges = strategy_state.get("edge_distribution", [])
    edge_stats = analyze_edge_distribution(edges)
    print(f"  Edge samples: {edge_stats.get('count', 0)}")

    # Generate directives
    print("\nGenerating research directives...")
    directives = generate_research_directives(
        edge_stats, paper_analysis, live_analysis, strategy_state,
    )
    print(f"  Generated {len(directives)} directives")

    # Build document
    print("\nBuilding document...")
    doc = build_document(
        strategy_state, paper_analysis, live_analysis, edge_stats, directives,
        intel_snapshot=intel_snapshot,
    )

    # Save
    doc.save(str(OUTPUT))
    print(f"\nSaved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")
    print("=" * 60)


if __name__ == "__main__":
    main()
