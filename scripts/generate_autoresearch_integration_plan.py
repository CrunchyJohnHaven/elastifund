#!/usr/bin/env python3
"""Generate a full .docx plan for integrating autoresearch learnings into Elastifund."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ops" / "AUTORESEARCH_INTEGRATION_PLAN.docx"


TITLE_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT_COLOR = RGBColor(0x1A, 0x7A, 0x3A)
BODY_COLOR = RGBColor(0x20, 0x20, 0x20)
MUTED_COLOR = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D9D9D9")
        borders.append(el)
    tbl_pr.append(borders)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY_COLOR

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = TITLE_COLOR


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = TITLE_COLOR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED_COLOR


def add_meta(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    row = table.rows[0]
    row.cells[0].text = "Field"
    row.cells[1].text = "Value"
    for cell in row.cells:
        set_cell_shading(cell, "EAF2F8")
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for key, value in items:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value

    doc.add_paragraph()


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(1)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(2)


def add_section(doc: Document, heading: str, paragraphs: list[str] | None = None, bullets: list[str] | None = None) -> None:
    doc.add_heading(heading, level=1)
    if paragraphs:
        for para in paragraphs:
            add_body(doc, para)
    if bullets:
        add_bullets(doc, bullets)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "DDEBF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    hdr = table.rows[0].cells
    for idx, label in enumerate(headers):
        hdr[idx].text = label
        set_cell_shading(hdr[idx], header_fill)
        for run in hdr[idx].paragraphs[0].runs:
            run.bold = True

    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    doc.add_paragraph()


def build_doc() -> None:
    doc = Document()
    style_doc(doc)
    today = date(2026, 3, 8)

    add_title(
        doc,
        "Autoresearch Integration Plan for Elastifund",
        "Execution plan for importing Andrej Karpathy's autoresearch patterns into the Elastifund research and flywheel stack",
    )
    add_meta(
        doc,
        [
            ("Date", today.isoformat()),
            ("Prepared by", "Codex"),
            ("Primary external reference", "github.com/karpathy/autoresearch"),
            ("Primary internal references", "research/karpathy_autoresearch_report.md, src/research_loop.py, flywheel/reporting.py"),
            ("Document purpose", "Turn the autoresearch comparison into a concrete build plan with milestones, repo changes, and governance boundaries"),
        ],
    )

    add_section(
        doc,
        "Executive Summary",
        paragraphs=[
            "Elastifund already shares the philosophy of autoresearch: iterate continuously, kill weak ideas quickly, and treat failures as first-class evidence. The missing piece is experimental geometry. Karpathy's system works because it reduces autonomous research to one mutable surface, one immutable evaluator, one fixed budget, one scalar objective, one append-only ledger, and one running-best graph.",
            "The right move is not to copy autoresearch across the whole repository. Elastifund has live-trading code, control-plane code, documentation, and public evidence artifacts. Whole-repo self-modification would mix research variance with operational risk. Instead, Elastifund should import the pattern lane-by-lane, starting with the forecast and calibration lane where the objective can be frozen cleanly.",
            "This plan proposes a four-week integration sequence beginning Monday, March 9, 2026. Week 1 freezes a benchmark contract. Week 2 builds a lane-local autoresearch harness. Week 3 runs bounded calibration experiments. Week 4 expands the same pattern into the strategy lane and connects the results to the flywheel control plane.",
        ],
    )

    add_section(
        doc,
        "Why This Matters Now",
        paragraphs=[
            "The repository now has a Karpathy comparison report and an autoresearch-style progress renderer. The current graph, generated from the existing run artifacts, shows 18 runs and only 1 kept high-water mark when using the current top-hypothesis composite score as a proxy objective. That is useful evidence: the repo has iteration logs, but it does not yet have a frozen benchmark lane strong enough to produce a trustworthy improvement frontier.",
            "Integrating autoresearch properly would sharpen three things at once: how the AI agent experiments, how the repo records evidence, and how the flywheel decides what deserves promotion. It would also create a cleaner context bundle for future deep research and external contributors.",
        ],
        bullets=[
            "It turns 'continuous research' into a measurable hill-climb.",
            "It separates benchmark wins from narrative wins.",
            "It gives the repo the same keep/discard graph discipline that makes autoresearch legible.",
            "It creates a bridge between experimentation and the existing flywheel control plane.",
        ],
    )

    add_section(
        doc,
        "What We Should Adopt from Autoresearch",
        bullets=[
            "One scalar objective per autonomous lane. The agent should optimize one score, not a bundle of mixed governance metrics.",
            "One immutable evaluator per lane. The harness must not drift while the agent is experimenting.",
            "One small mutable surface per lane. Start with a single calibrator or a single strategy module, not the whole repo.",
            "An append-only results ledger with keep, discard, and crash semantics.",
            "A running-best progress graph exported after every benchmark batch.",
            "A human-authored lane contract that says what is in scope, what is out of scope, what command runs the benchmark, and what qualifies as a keep.",
        ],
    )

    add_section(
        doc,
        "What We Should Not Copy Literally",
        bullets=[
            "Do not allow repo-wide self-modification. Live execution, deployment, and public reporting code must stay outside the first benchmark loops.",
            "Do not use one metric for the entire business. Forecast quality, strategy quality, and execution quality each need their own benchmark lane.",
            "Do not let benchmark wins auto-promote code into live capital. Benchmark lanes feed the flywheel; they do not override it.",
            "Do not change scoring weights or data windows during an experiment wave. Moving the ruler invalidates the frontier.",
        ],
    )

    doc.add_heading("Autoresearch Pattern Mapping", level=1)
    add_table(
        doc,
        ["Autoresearch Pattern", "Why It Works", "Elastifund Adaptation", "Primary Repo Surface"],
        [
            [
                "One mutable file",
                "Keeps the search space small and diffs reviewable",
                "Start with one calibrator file or one strategy module per lane",
                "bot/adaptive_platt.py or one file in src/strategies/",
            ],
            [
                "Immutable evaluator",
                "Makes runs comparable across time",
                "Create benchmark packages with fixed data, splits, and cost assumptions",
                "benchmarks/calibration_v1/, benchmarks/strategy_v1/",
            ],
            [
                "Single scalar objective",
                "Removes ambiguity in keep/discard decisions",
                "Use benchmark_score derived from held-out Brier and ECE for calibration lane",
                "scripts/run_calibration_benchmark.py",
            ],
            [
                "results.tsv ledger",
                "Creates an auditable experiment history",
                "Write lane-local ledgers under versioned research results paths",
                "research/results/calibration/results.tsv",
            ],
            [
                "Running-best graph",
                "Makes progress visible at a glance",
                "Promote progress graph exports into repo-tracked lane artifacts",
                "research/autoresearch_progress.svg and lane-specific successors",
            ],
            [
                "program.md contract",
                "Gives the agent a stable operating system",
                "Create lane contracts for calibration and strategy search",
                "research/programs/calibration_lane.md",
            ],
        ],
    )

    add_section(
        doc,
        "Target Operating Model",
        paragraphs=[
            "Elastifund should run three distinct improvement loops. Only the first loop needs to be built immediately. The other two inherit the same pattern after the first loop is stable.",
        ],
    )
    add_table(
        doc,
        ["Lane", "Optimization Target", "Mutable Surface", "Immutable Harness", "Promotion Path"],
        [
            [
                "Calibration lane",
                "Held-out benchmark_score built from Brier and ECE",
                "bot/adaptive_platt.py",
                "Frozen resolved-market slice and evaluator",
                "If benchmark improves, merge into research branch and replay in paper calibration",
            ],
            [
                "Strategy lane",
                "Fixed replay score on frozen market slice",
                "Single strategy module plus config",
                "Frozen backtest package and cost model",
                "If benchmark improves, enter paper or shadow queue",
            ],
            [
                "Execution lane",
                "Fill-adjusted edge and slippage fidelity on replay or shadow data",
                "Execution logic module only",
                "Replay or shadow evaluation harness",
                "Only after paper or shadow evidence and explicit flywheel approval",
            ],
        ],
        header_fill="E2F0D9",
    )

    add_section(
        doc,
        "Governance Guardrails",
        bullets=[
            "No benchmark lane may directly modify bot/jj_live.py as part of the first implementation wave.",
            "No benchmark win may bypass docs/ops/Flywheel_Control_Plane.md promotion boundaries.",
            "No experiment may rewrite the evaluator, data manifest, or acceptance criteria during an active benchmark wave.",
            "Every lane result must remain reproducible from one command and one manifest.",
            "Every public claim must preserve the repo's current honesty standard: backtest, replay, paper, and live must remain explicitly labeled.",
        ],
    )

    doc.add_heading("Four-Week Integration Plan", level=1)
    add_body(doc, "Schedule anchor: Monday, March 9, 2026 through Friday, April 3, 2026.")

    phases = [
        {
            "title": "Phase 0 — Freeze Benchmark v1",
            "dates": "March 9-10, 2026",
            "goal": "Create one immutable calibration benchmark contract so the repo has a stable ruler before autonomous iteration begins.",
            "deliverables": [
                "benchmarks/calibration_v1/README.md describing scope, dataset, splits, and objective",
                "benchmarks/calibration_v1/manifest.json with dataset snapshot, date cutoffs, and artifact checksums",
                "scripts/run_calibration_benchmark.py producing benchmark_score, brier, ece, and metadata",
                "research/programs/calibration_lane.md defining in-scope files, out-of-scope files, keep/discard criteria, and safety boundaries",
                "research/results/calibration/results.tsv initialized with a baseline row",
            ],
            "exit": [
                "A clean machine can run one command and obtain the same baseline benchmark result.",
                "The objective function is frozen for the full first experiment wave.",
                "The benchmark package is reviewable and versioned in git.",
            ],
        },
        {
            "title": "Phase 1 — Build the Lane-Local Autoresearch Harness",
            "dates": "March 11-13, 2026",
            "goal": "Replicate the core autoresearch mechanics inside the calibration lane without touching live execution.",
            "deliverables": [
                "scripts/run_lane_autoresearch.py that edits or evaluates one mutable surface and records keep/discard/crash decisions",
                "research/results/calibration/progress.svg and progress.tsv exported from the lane ledger",
                "Branch naming convention such as codex/autoresearch-calibration-mar2026",
                "A benchmark packet format containing commit, objective, brier, ece, status, and short description",
                "A README note explaining that calibration lane wins do not equal live-trading wins",
            ],
            "exit": [
                "The agent can run baseline, mutate, benchmark, log, and revert or keep with no manual bookkeeping.",
                "Progress is visible as a running-best graph after each batch.",
                "The harness can crash safely without corrupting prior results.",
            ],
        },
        {
            "title": "Phase 2 — Run the First Calibration Experiment Wave",
            "dates": "March 16-20, 2026",
            "goal": "Use the harness to search a bounded family of calibration improvements on a fixed benchmark slice.",
            "deliverables": [
                "25-50 logged calibration experiments with keep/discard/crash outcomes",
                "A short experiment taxonomy: rolling-window size, isotonic vs Platt fallback, category shrinkage, disagreement-aware shrinkage, and base-rate priors",
                "A ranked shortlist of retained changes with exact benchmark deltas",
                "One merge-ready patch if the benchmark frontier improves materially",
            ],
            "exit": [
                "At least one experiment wave is complete and reproducible.",
                "If no change wins, that is written down as an honest null result.",
                "If a change wins, it is replayed outside the mutation loop before any broader adoption.",
            ],
        },
        {
            "title": "Phase 3 — Expand to the Strategy Lane and Connect to the Flywheel",
            "dates": "March 23-April 3, 2026",
            "goal": "Apply the same pattern to a frozen strategy benchmark and wire benchmark outputs into the existing control plane.",
            "deliverables": [
                "benchmarks/strategy_v1/ package with fixed replay slice and cost assumptions",
                "research/programs/strategy_lane.md defining one strategy-family mutation surface at a time",
                "A lane-to-flywheel bridge so benchmark wins generate review tasks rather than silent code drift",
                "Website or docs-ready benchmark notes describing what the graph measures and what it does not measure",
            ],
            "exit": [
                "Calibration lane artifacts feed the flywheel task queue.",
                "Strategy lane experiments use the same keep/discard discipline.",
                "The repo has a stable pattern that can be reused for later execution-lane work.",
            ],
        },
    ]

    for phase in phases:
        doc.add_heading(phase["title"], level=2)
        add_body(doc, f"Date range: {phase['dates']}")
        add_body(doc, f"Goal: {phase['goal']}")
        doc.add_heading("Deliverables", level=3)
        add_bullets(doc, phase["deliverables"])
        doc.add_heading("Exit Criteria", level=3)
        add_bullets(doc, phase["exit"])

    doc.add_heading("Immediate Repo Additions Recommended", level=1)
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["benchmarks/calibration_v1/", "Immutable benchmark package for the first lane"],
            ["research/programs/calibration_lane.md", "Lane contract equivalent to program.md"],
            ["research/results/calibration/results.tsv", "Append-only experiment ledger"],
            ["research/results/calibration/progress.svg", "Versioned running-best graph"],
            ["scripts/run_calibration_benchmark.py", "Immutable evaluator entrypoint"],
            ["scripts/run_lane_autoresearch.py", "Autonomous keep/discard loop for one lane"],
            ["benchmarks/strategy_v1/", "Second benchmark lane after calibration"],
            ["research/programs/strategy_lane.md", "Strategy-lane contract"],
        ],
    )

    add_section(
        doc,
        "Detailed Workstreams",
        paragraphs=[
            "The plan is easier to execute if it is split into workstreams owned by function rather than by file. The same person can own several workstreams, but the work itself should be tracked separately.",
        ],
    )
    add_table(
        doc,
        ["Workstream", "Primary Outcome", "Key Tasks", "Definition of Done"],
        [
            [
                "Benchmark design",
                "Frozen benchmark contract",
                "Snapshot data, define objective, define manifests, define command",
                "One command reproduces the baseline exactly",
            ],
            [
                "Lane contracts",
                "Clear agent operating system",
                "Write calibration_lane.md and strategy_lane.md",
                "No ambiguity about mutable scope or keep rules",
            ],
            [
                "Automation harness",
                "Autonomous experiment loop",
                "Implement run, log, graph, revert, and keep behavior",
                "The loop can run repeatedly without manual bookkeeping",
            ],
            [
                "Reporting",
                "Human-readable and machine-readable artifacts",
                "Emit results.tsv, progress.svg, summary.md",
                "Artifacts support both local review and public explanation",
            ],
            [
                "Flywheel integration",
                "Benchmark outcomes become review tasks",
                "Bridge benchmark wins and null results into flywheel outputs",
                "No benchmark result is invisible to the control plane",
            ],
            [
                "Governance",
                "Safety and evidence discipline",
                "Preserve labels, guardrails, and merge review checkpoints",
                "No silent widening of live risk",
            ],
        ],
        header_fill="FCE4D6",
    )

    add_section(
        doc,
        "Objective Function Recommendation for the First Lane",
        paragraphs=[
            "The first lane should be the forecast and calibration lane. It already has the cleanest objective candidates in the repo. The benchmark should produce one scalar called benchmark_score so the agent knows exactly what to optimize.",
            "Recommended v1 formula: benchmark_score = -(brier + 0.25 * ece). This makes lower Brier and lower calibration error both help, while keeping Brier dominant. The exact coefficient can be tuned once before the benchmark is frozen, but it must not change during an active experiment wave.",
            "Secondary metrics such as log loss, per-category Brier, and calibration plots should still be recorded. They are diagnostic outputs, not optimization targets.",
        ],
        bullets=[
            "Primary metric: benchmark_score",
            "Diagnostic metrics: brier, ece, log_loss, per-category brier, confidence-band drift",
            "Tie-break rule: prefer the simpler implementation when benchmark_score is effectively unchanged",
        ],
    )

    doc.add_heading("Success Metrics", level=1)
    add_table(
        doc,
        ["Category", "Metric", "Success Threshold", "Why It Matters"],
        [
            ["Benchmark stability", "Baseline rerun variance", "Within predefined tolerance band", "Proves the evaluator is stable enough to trust"],
            ["Experiment throughput", "Logged experiments per wave", "25 or more in first calibration wave", "Shows the harness is actually usable"],
            ["Benchmark quality", "Number of retained improvements", "At least one material win or a documented null result", "Avoids vague claims of progress"],
            ["Reporting quality", "Progress artifact freshness", "Graph and ledger updated after each wave", "Makes the frontier legible"],
            ["Governance quality", "Unauthorized live-surface mutations", "Zero", "Prevents research convenience from widening risk"],
            ["Flywheel integration", "Benchmark-to-task propagation", "All retained improvements become explicit tasks or review packets", "Connects experimentation to deployment decisions"],
        ],
        header_fill="E2F0D9",
    )

    add_section(
        doc,
        "Risk Register and Mitigations",
        bullets=[
            "Risk: the benchmark contract is too loose and the graph becomes noise. Mitigation: freeze a small but reviewable dataset and document checksums.",
            "Risk: the mutable surface is too broad and experiment interpretation collapses. Mitigation: one file or one narrow module per lane.",
            "Risk: benchmark wins do not transfer to paper or shadow behavior. Mitigation: every retained change must be replayed outside the autonomous loop before adoption.",
            "Risk: the team confuses benchmark progress with live profitability. Mitigation: keep labeling boundaries explicit in docs and artifacts.",
            "Risk: agents spend time improving the harness instead of the benchmark target. Mitigation: evaluator and manifests remain read-only inside an experiment wave.",
        ],
    )

    add_section(
        doc,
        "Roles and Operating Responsibilities",
        bullets=[
            "Human owner: freeze benchmark versions, approve lane boundaries, review retained changes, and decide what merges.",
            "Coding agent: mutate only the approved lane surface, run benchmarks, log results, generate progress artifacts, and prepare review packets.",
            "Flywheel control plane: record retained changes as tasks, recommendations, or review artifacts; never as automatic live promotions.",
        ],
    )

    add_section(
        doc,
        "Deep Research Handoff Bundle",
        paragraphs=[
            "If this plan is handed to a deeper research system, the context bundle should include both the external autoresearch materials and the internal Elastifund files that define current architecture, reporting, and calibration logic.",
        ],
        bullets=[
            "External: karpathy/autoresearch README.md, program.md, train.py, prepare.py, progress.png",
            "Internal: research/karpathy_autoresearch_report.md",
            "Internal: src/research_loop.py, src/hypothesis_manager.py, src/reporting.py, src/edge_registry.py",
            "Internal: bot/adaptive_platt.py and src/confidence_calibration.py",
            "Internal: docs/ARCHITECTURE.md, docs/PERFORMANCE.md, docs/RESEARCH_LOG.md, docs/website/benchmark-methodology.md, docs/ops/Flywheel_Control_Plane.md",
            "Internal artifacts: research/autoresearch_progress.tsv and research/autoresearch_progress.svg",
        ],
    )

    doc.add_heading("First 10 Tasks to Open Immediately", level=1)
    add_numbers(
        doc,
        [
            "Create benchmarks/calibration_v1/ with manifest, README, and frozen split definition.",
            "Implement scripts/run_calibration_benchmark.py with benchmark_score, brier, and ece outputs.",
            "Write research/programs/calibration_lane.md with keep/discard/crash rules.",
            "Initialize research/results/calibration/results.tsv with a baseline row.",
            "Clone the current progress renderer into a lane-specific exporter that reads the calibration ledger.",
            "Constrain the first mutable surface to bot/adaptive_platt.py only.",
            "Define the first 5-7 calibration experiment families before opening the loop.",
            "Wire retained benchmark wins into flywheel task generation rather than silent code adoption.",
            "Add a short benchmark disclaimer section to docs/website/benchmark-methodology.md or a sibling page.",
            "Schedule the first experiment wave for March 16-20, 2026 and review results at the end of that week.",
        ],
    )

    add_section(
        doc,
        "Recommended Decision",
        paragraphs=[
            "Proceed with the four-week lane-by-lane integration, starting Monday, March 9, 2026 with the calibration benchmark contract. Do not widen scope to live execution during this first wave. The right success condition is not 'the AI rewrote the whole repo.' The right success condition is that Elastifund can point to a frozen benchmark lane, a keep/discard ledger, and a running-best graph that honestly shows whether the system is getting better.",
        ],
    )

    doc.sections[-1].start_type = WD_SECTION.NEW_PAGE
    doc.add_heading("Appendix: Current-State Anchor", level=1)
    add_body(doc, "As of March 8, 2026, the repo already contains a Karpathy comparison report and an autoresearch-style progress export. The current export is intentionally limited: it reads existing reports/run_*_metrics.json files and uses top-hypothesis composite score as a proxy objective. It shows 18 runs and only 1 kept high-water mark. This is the proof point for why a frozen lane benchmark is the correct next build, not a reason to abandon the pattern.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build_doc()
