#!/usr/bin/env python3
"""
Generate 6 deep research prompt .docx files for the Polymarket quant project.
Uses python-docx exclusively. Self-contained single script.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def add_section_heading(doc, text):
    """Add a bold section heading with a bit of spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_body(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    style = p.style
    font = style.font
    font.size = Pt(11)


def add_bullet(doc, text):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def build_doc(filename, title, objective, prompt_text, expected_outputs, data_sources):
    """Build a single .docx with standardised structure."""
    doc = Document()

    # -- Title --
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -- Research Objective --
    add_section_heading(doc, "Research Objective")
    add_body(doc, objective)

    # -- Prompt --
    add_section_heading(doc, "Prompt")
    # Prompt may have multiple paragraphs separated by \n\n
    for block in prompt_text.strip().split("\n\n"):
        add_body(doc, block.strip())

    # -- Expected Outputs --
    add_section_heading(doc, "Expected Outputs")
    for item in expected_outputs:
        add_bullet(doc, item)

    # -- Data Sources --
    add_section_heading(doc, "Data Sources")
    for item in data_sources:
        add_bullet(doc, item)

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"  Created: {path}")


# ============================================================================
# 1. Claude Calibration Analysis
# ============================================================================
build_doc(
    filename="01_claude_calibration_analysis.docx",
    title="Claude Calibration Analysis \u2014 Probability Estimation Accuracy on Prediction Markets",
    objective=(
        "Determine how well Claude Haiku estimates probabilities on resolved prediction markets. "
        "Quantify calibration quality, identify systematic biases (overconfidence, anchoring to market "
        "prices, category-specific errors), and establish whether Claude's estimates add alpha over "
        "naively trusting market prices."
    ),
    prompt_text="""\
Perform a systematic calibration analysis of Claude Haiku's probability estimation on prediction markets using the following protocol:

1. DATASET CONSTRUCTION
Select 100 resolved Polymarket markets spanning at least five categories: U.S. politics, international geopolitics, cryptocurrency price targets, sports outcomes, and weather/climate events. For each market, record: (a) the market question exactly as phrased, (b) the resolution date, (c) the final outcome (YES/NO), and (d) the market mid-price at three snapshots \u2014 30 days before resolution, 7 days before resolution, and 24 hours before resolution.

2. BLIND ESTIMATION (No Market Price)
For each market, present Claude Haiku with the question text and the current date context ONLY. Do not reveal the market price. Ask Claude to output a probability between 0.01 and 0.99. Record this as P_blind.

3. ANCHORED ESTIMATION (Market Price Shown)
Repeat the estimation but now include the market mid-price in the prompt. Record this as P_anchored.

4. CALIBRATION METRICS
Compute the following for both P_blind and P_anchored:
- Brier score (mean squared error vs. binary outcome)
- Calibration curve: bin predictions into deciles (0.0\u20130.1, 0.1\u20130.2, \u2026, 0.9\u20131.0) and plot predicted probability vs. observed frequency
- Calibration error (ECE): expected calibration error across bins
- Resolution component and reliability component of the Brier decomposition
- Log-loss (cross-entropy)

5. BIAS ANALYSIS
- Overconfidence: compare average confidence (distance from 0.5) to average accuracy
- Anchoring: regress (P_anchored \u2212 P_blind) on the market price to quantify anchoring strength
- Category breakdown: compute per-category Brier scores and identify which categories Claude is best/worst at
- Temporal analysis: does Claude's accuracy degrade the further out from resolution?

6. COMPARISON BASELINES
- Baseline 1: always predict 0.5 (maximum ignorance)
- Baseline 2: always predict the market price (efficient market hypothesis)
- Baseline 3: a simple 50/50 blend of Claude + market price

Determine whether Claude adds statistically significant information beyond the market price (paired t-test on squared errors, p < 0.05).\
""",
    expected_outputs=[
        "Brier score table: P_blind, P_anchored, market-price baseline, and 50/50 blend across all 100 markets",
        "Calibration curve plots (or data tables) for P_blind and P_anchored",
        "Anchoring regression coefficient with confidence interval",
        "Per-category Brier score breakdown (politics, crypto, sports, weather, geopolitics)",
        "Statistical significance test result (paired t-test) for Claude vs. market baseline",
        "Ranked list of market categories by Claude's relative edge or deficit",
        "Concrete recommendations: when to trust Claude's estimate, when to defer to the market",
    ],
    data_sources=[
        "Polymarket Gamma API \u2014 resolved markets endpoint (https://gamma-api.polymarket.com/markets?closed=true)",
        "Polymarket CLOB API \u2014 historical price snapshots",
        "Claude Haiku API (Anthropic) \u2014 for generating probability estimates",
        "Python libraries: numpy, scipy.stats, sklearn.calibration, matplotlib",
    ],
)

# ============================================================================
# 2. Market Microstructure Deep Dive
# ============================================================================
build_doc(
    filename="02_market_microstructure_deep_dive.docx",
    title="Polymarket Market Microstructure Deep Dive",
    objective=(
        "Develop a comprehensive understanding of Polymarket's Central Limit Order Book (CLOB) "
        "architecture, fee structure, NegRisk framework, liquidity dynamics, and the Builder program's "
        "impact on order flow \u2014 all of which directly affect strategy design and execution quality."
    ),
    prompt_text="""\
Conduct a thorough investigation of Polymarket's market microstructure covering every layer from on-chain settlement to API-level order management. Structure the research as follows:

1. CLOB ARCHITECTURE
- Explain how Polymarket's off-chain CLOB works: order matching engine, order types (limit, market, FOK, GTC, GTD), partial fills, and self-trade prevention.
- Describe the relationship between the off-chain order book and on-chain CTF (Conditional Token Framework) settlement on Polygon.
- How are NegRisk markets structured? Explain the NegRisk adapter contract, how multi-outcome markets (e.g., "Who will win the 2024 election?") are decomposed into binary pairs, and how complement pricing works.
- What are the implications of NegRisk for arbitrage between YES and NO tokens across outcomes?

2. FEE STRUCTURE
- Detail the current maker/taker fee schedule. Are there volume-based tiers?
- How do fees interact with the spread? What is the minimum profitable spread for a maker?
- Explain the reward/rebate programs for liquidity providers.

3. LIQUIDITY ANALYSIS
- Using the Gamma API and CLOB API, analyze bid-ask spreads for the top 20 markets by volume. Report: median spread, 95th percentile spread, spread as a function of market price (are spreads wider near 0.50?).
- Analyze liquidity depth: how much capital is needed to move the price by 1%, 5%, 10%?
- Document liquidity patterns by time of day (UTC). Is there a consistent intraday pattern? Are spreads wider overnight?
- How does liquidity behave around major news events? Document spread blow-outs and recovery times.

4. BUILDER PROGRAM AND GASLESS TRANSACTIONS
- Explain the Builder/Operator model: how do approved operators submit orders on behalf of users without gas fees?
- What are the technical requirements (API keys, POLY_PROXY_WALLET, signature types POLY_GNOSIS_SAFE vs EOA)?
- How does gasless order flow differ from direct user orders in terms of latency, fill rates, and order size distribution?
- What fraction of total volume flows through Builders vs. direct API users?

5. ORDER FLOW TOXICITY
- Is there evidence of informed order flow (e.g., large orders consistently preceding news events)?
- Estimate the VPIN (Volume-Synchronized Probability of Informed Trading) for high-volume political markets.
- How should a small ($75 bankroll) algorithmic trader position orders to minimize adverse selection?\
""",
    expected_outputs=[
        "Architecture diagram of CLOB + on-chain settlement flow",
        "Fee schedule table with break-even spread calculations",
        "Bid-ask spread statistics (median, p5, p95) for top 20 markets",
        "Liquidity depth chart: price impact as a function of order size",
        "Intraday spread/volume heatmap by hour (UTC)",
        "NegRisk arbitrage opportunity analysis with worked example",
        "Builder program integration checklist for bot development",
        "Adverse selection risk assessment and mitigation strategies for small bankrolls",
    ],
    data_sources=[
        "Polymarket CLOB API documentation (https://docs.polymarket.com/)",
        "Polymarket Gamma API \u2014 markets, orders, trades endpoints",
        "Polygon blockchain explorer (Polygonscan) \u2014 CTF contract interactions",
        "NegRisk adapter contract source code on GitHub",
        "Academic literature: Kyle (1985), Glosten-Milgrom (1985) for microstructure theory",
    ],
)

# ============================================================================
# 3. Weather Arbitrage Optimization
# ============================================================================
build_doc(
    filename="03_weather_arbitrage_optimization.docx",
    title="NOAA Weather Arbitrage Optimization for Polymarket",
    objective=(
        "Maximize the edge and execution quality of the weather-based arbitrage strategy by deeply "
        "understanding NOAA forecast accuracy, market update latency, optimal timing, and the value "
        "of combining multiple numerical weather prediction models."
    ),
    prompt_text="""\
Conduct a comprehensive analysis to optimize the NOAA weather arbitrage strategy on Polymarket. This strategy exploits the lag between official NOAA forecast releases and Polymarket weather market price updates.

1. NOAA FORECAST ACCURACY ANALYSIS
- For the NWS point forecast API (api.weather.gov), quantify forecast accuracy for:
  * Daily high/low temperature: MAE and bias by lead time (1-day, 2-day, 3-day, 5-day, 7-day)
  * Probability of precipitation (PoP): calibration curve by lead time
  * Precipitation amount: MAE by lead time and climate region
- Break down accuracy by city (focus on cities with active Polymarket weather markets: NYC, LA, Chicago, Miami, Denver, Phoenix).
- Identify systematic biases: does NOAA consistently over- or under-forecast temperature extremes? Is PoP systematically overconfident?
- How does forecast accuracy change seasonally? Are winter forecasts less accurate than summer?

2. MARKET UPDATE LATENCY
- Document the exact schedule of NOAA forecast issuance: when do zone forecasts, point forecasts, and model runs (GFS, NAM) become available?
- After NOAA issues an updated forecast, how quickly do Polymarket weather markets adjust? Measure this by:
  * Polling market prices every 60 seconds around known forecast release times
  * Computing the time delta between forecast release and when the market price moves more than 2% toward the forecast-implied probability
- Is the lag consistent, or does it vary by market liquidity and time of day?

3. OPTIMAL BET TIMING
- Given the forecast release schedule and observed market lag, what is the optimal window for placing bets?
- Should we place limit orders in advance or market orders immediately after forecast release?
- How does order size affect fill probability and slippage during the post-forecast window?
- Model the tradeoff: betting earlier captures more edge but on less accurate forecasts; betting later uses more accurate forecasts but captures less mispricing.

4. MULTI-MODEL ENSEMBLE
- Compare forecast accuracy of: GFS (NOAA), ECMWF (European), NAM (regional), HRRR (high-resolution rapid refresh), and the NWS blended forecast.
- Does a simple ensemble (average of GFS + ECMWF) outperform any single model for Polymarket-relevant variables?
- For temperature threshold markets ("Will NYC high exceed 80\u00b0F?"), how should we convert continuous forecasts to binary probabilities? Analyze: (a) historical forecast error distribution \u2192 CDF approach, (b) ensemble spread as uncertainty estimate.
- Is ECMWF data worth paying for given the marginal accuracy improvement?

5. HISTORICAL PRICING ANALYSIS
- Collect historical Polymarket weather market data for the past 6 months.
- For resolved markets: what was the average mispricing relative to NOAA forecasts? How did mispricing evolve as resolution approached?
- Estimate the realistic edge (after fees and slippage) for the weather strategy.
- How many weather markets are active simultaneously, and what is the total addressable capital deployment?\
""",
    expected_outputs=[
        "NOAA forecast accuracy table: MAE/bias by city, variable, and lead time",
        "Calibration curve for NOAA PoP forecasts vs. observed precipitation",
        "Market update latency distribution (minutes after forecast release)",
        "Optimal betting window recommendation with expected edge decay curve",
        "Multi-model ensemble accuracy comparison table",
        "Binary probability conversion methodology for threshold markets",
        "Historical mispricing analysis: average edge by market type and timing",
        "Capital deployment plan: how many simultaneous weather bets, position sizing per bet",
    ],
    data_sources=[
        "NOAA/NWS API (api.weather.gov) \u2014 point forecasts, zone forecasts, observations",
        "NOAA GFS model data via NOMADS (nomads.ncep.noaa.gov)",
        "ECMWF open data (if available) or ECMWF API (paid tier)",
        "Polymarket Gamma API \u2014 weather market historical prices",
        "Iowa Environmental Mesonet (IEM) \u2014 historical observations for verification",
        "Academic literature: NWS forecast verification reports, ensemble calibration methods",
    ],
)

# ============================================================================
# 4. Kelly Criterion for Prediction Markets
# ============================================================================
build_doc(
    filename="04_kelly_criterion_for_prediction_markets.docx",
    title="Kelly Criterion for Correlated Binary Prediction Markets",
    objective=(
        "Derive and validate optimal position sizing rules for a $75-bankroll prediction market "
        "bot operating across correlated binary markets, balancing growth rate maximization against "
        "risk of ruin in a setting with uncertain edge estimates."
    ),
    prompt_text="""\
Conduct deep research on applying the Kelly criterion and its variants to position sizing in correlated binary prediction markets. This is for a bot with $75 starting capital trading on Polymarket.

1. KELLY CRITERION FUNDAMENTALS FOR BINARY MARKETS
- Derive the Kelly fraction for a single binary prediction market bet: f* = (p\u00b7b - q) / b, where p = true probability, q = 1-p, b = net odds (payout ratio).
- On Polymarket, if the market price is 0.40 and we believe the true probability is 0.55, what is the Kelly bet? Work through the exact calculation including fees.
- Explain why full Kelly is theoretically optimal (maximizes log-utility / geometric growth rate) but practically dangerous.

2. FRACTIONAL KELLY AND UNCERTAIN EDGE
- Derive the relationship between Kelly fraction and: (a) growth rate, (b) variance of returns, (c) probability of drawdown exceeding X%.
- For our observed edge distribution (mean edge = 25.7%, range 5%\u201347% across different market categories), what is the optimal fixed fractional Kelly multiplier?
- Model edge uncertainty: if our estimated edge has a standard error of \u00b110%, how does this affect optimal sizing? Derive the "shrunk Kelly" formula that accounts for estimation error.
- Simulate 1000 trajectories of 200 bets each at full Kelly, half Kelly, quarter Kelly, and tenth Kelly. For each, report: median terminal wealth, probability of 50% drawdown, probability of ruin (bankroll < $5), and Sharpe ratio of returns.

3. CORRELATED POSITIONS
- We often have multiple simultaneous bets on related outcomes (e.g., 5 GTA VI release date markets at different thresholds, or multiple political markets affected by the same election). Derive the multi-asset Kelly criterion for correlated binary bets.
- Provide the matrix formulation: f* = \u03a3\u207b\u00b9 \u00b7 \u03bc, where \u03a3 is the covariance matrix of returns and \u03bc is the expected excess return vector. How do we estimate \u03a3 for prediction markets?
- Worked example: we hold YES positions in "GTA VI releases in 2025 Q4" (edge = 30%), "GTA VI releases before 2026" (edge = 20%), and "Take-Two stock above $200 by Dec 2025" (edge = 15%). These are highly correlated. What are the optimal position sizes?
- How does ignoring correlation lead to over-betting? Quantify the excess risk.

4. RISK OF RUIN AT $75 BANKROLL
- With a $75 starting bankroll and the observed edge/odds distribution, compute the probability of ruin (bankroll falls below $1) under different Kelly fractions.
- What is the minimum bankroll needed for the bot to have <5% ruin probability over 500 bets at half Kelly?
- Derive the expected time (number of bets) to grow from $75 to $500 under quarter Kelly, given the observed edge distribution.
- How should bet sizing change as the bankroll grows? Define bankroll thresholds for adjusting the Kelly fraction upward.

5. PRACTICAL IMPLEMENTATION
- Polymarket has minimum order sizes and discrete price increments. How do we round Kelly-optimal bets to feasible order sizes?
- When the Kelly-optimal bet is less than the minimum order size, should we skip the bet or bet the minimum?
- How frequently should we recalculate position sizes as the bankroll changes intra-day?
- Implement a simple Python function: kelly_bet(bankroll, edge, market_price, kelly_fraction, fee_rate) \u2192 position_size_usd.\
""",
    expected_outputs=[
        "Derivation of Kelly fraction for Polymarket binary bets including fee adjustment",
        "Simulation results table: terminal wealth distribution at full/half/quarter/tenth Kelly over 200 bets",
        "Ruin probability curves as a function of Kelly fraction for $75 bankroll",
        "Multi-asset Kelly formula with worked example for correlated GTA VI markets",
        "Optimal Kelly fraction recommendation given observed edge uncertainty (\u00b110%)",
        "Bankroll growth trajectory chart: expected path from $75 to $500",
        "Python implementation of kelly_bet() function with fee and minimum order handling",
        "Decision rules: when to skip bets, when to adjust Kelly fraction, bankroll thresholds",
    ],
    data_sources=[
        "Kelly (1956) \u2014 original paper: 'A New Interpretation of Information Rate'",
        "Thorp (2006) \u2014 'The Kelly Criterion in Blackjack, Sports Betting, and the Stock Market'",
        "MacLean, Thorp, Ziemba (2011) \u2014 'The Kelly Capital Growth Investment Criterion'",
        "Polymarket fee schedule and minimum order documentation",
        "Bot's historical trade log \u2014 observed edge distribution (mean 25.7%, range 5%\u201347%)",
        "Monte Carlo simulation (Python: numpy, scipy)",
    ],
)

# ============================================================================
# 5. Political Event Prediction Edge
# ============================================================================
build_doc(
    filename="05_political_event_prediction_edge.docx",
    title="LLM Edge on Political and Geopolitical Prediction Markets",
    objective=(
        "Rigorously evaluate whether large language models (specifically Claude) possess genuine, "
        "exploitable forecasting edge on political and geopolitical prediction markets, and identify "
        "the specific conditions and prompt strategies that maximize this edge."
    ),
    prompt_text="""\
Investigate whether LLMs have a real, persistent informational or analytical edge over prediction market prices for political and geopolitical events. This research directly informs whether the bot should trade political markets or restrict itself to data-driven categories (weather, sports statistics).

1. LLM vs. SUPERFORECASTER vs. MARKET ACCURACY
- Survey the academic literature comparing LLM forecasting to human superforecasters and prediction markets:
  * Halawi et al. (2024) \u2014 "Approaching Human-Level Forecasting with Language Models": what were the key findings? How close did GPT-4/Claude get to human superforecasters on Metaculus?
  * Schoenegger et al. (2024) \u2014 LLM performance on the Good Judgment Open tournament
  * Any other relevant papers from 2024-2025 on LLM forecasting accuracy
- Summarize: on which question types do LLMs match or beat markets? On which do they consistently underperform?

2. INFORMATION ASYMMETRY ANALYSIS
- What information does Claude have access to that might not be fully priced into prediction markets?
  * Training data cutoff implications: Claude knows historical base rates, political science research, and geopolitical analysis up to its training cutoff
  * Claude CANNOT access real-time news, polls, or insider information
- Conversely, what do markets have that Claude lacks?
  * Aggregated wisdom of crowds including domain experts
  * Real-time information flow (breaking news, leaked documents, insider trading)
  * Financial incentives that motivate careful analysis
- Given this asymmetry, in what specific scenarios might Claude have edge? Hypothesize and test.

3. PROMPT ENGINEERING FOR CALIBRATION
- Test the following prompt strategies on a set of 50 resolved political/geopolitical markets and measure Brier score for each:
  * Strategy A \u2014 Naive: "What is the probability that [event]?" (single-shot)
  * Strategy B \u2014 Chain-of-Thought: "Think step by step. Consider the key factors, historical precedents, and current conditions. Then estimate the probability."
  * Strategy C \u2014 Reference Class Forecasting: "First identify the reference class for this event. What is the base rate? Then adjust for specific factors."
  * Strategy D \u2014 Adversarial Debiasing: "List 3 reasons this will happen and 3 reasons it won't. Weight each reason. Then give your probability."
  * Strategy E \u2014 Superforecaster Persona: "You are a superforecaster with a track record of well-calibrated predictions. Use the techniques from Tetlock's 'Superforecasting': outside view first, then inside view adjustment, be willing to use precise probabilities, update incrementally."
  * Strategy F \u2014 Ensemble: Average of strategies B through E.
- Report Brier scores, calibration curves, and statistical significance of differences between strategies.

4. CATEGORY-SPECIFIC EDGE ANALYSIS
- Break down LLM accuracy by political subcategory:
  * U.S. presidential/congressional elections
  * International elections (UK, France, etc.)
  * Geopolitical events (wars, sanctions, treaties)
  * Policy decisions (Fed rate decisions, legislation passage)
  * Supreme Court decisions
  * Political appointment confirmations
- For each subcategory: does Claude beat the market, match it, or underperform? What is the magnitude of edge or deficit?

5. PRACTICAL TRADING IMPLICATIONS
- Given the findings above, should the bot trade political markets? Under what conditions?
- If yes: which subcategories, at what minimum edge threshold, with what Kelly fraction (accounting for higher uncertainty)?
- If no: is there value in using Claude's political forecasts as a secondary signal (e.g., confirming/rejecting trades identified by other strategies)?
- How should the bot handle the training data cutoff problem for events that depend heavily on recent developments?\
""",
    expected_outputs=[
        "Literature review summary table: LLM vs. superforecaster vs. market accuracy across studies",
        "Information asymmetry matrix: what Claude knows vs. what markets know",
        "Brier score comparison table across 6 prompt strategies on 50 resolved markets",
        "Best prompt strategy recommendation with statistical justification",
        "Category-specific edge analysis: Brier score by political subcategory",
        "Go/no-go recommendation for political market trading with specific conditions",
        "Prompt template for the bot to use when evaluating political markets",
        "Risk factors and failure modes for LLM-based political trading",
    ],
    data_sources=[
        "Metaculus \u2014 resolved political/geopolitical questions with community predictions",
        "Good Judgment Open \u2014 historical forecasting tournament data",
        "Polymarket \u2014 resolved political markets with historical price data",
        "Halawi et al. (2024) and related LLM forecasting papers",
        "Tetlock (2015) \u2014 'Superforecasting' methodology",
        "Claude API (Anthropic) \u2014 for running prompt strategy experiments",
    ],
)

# ============================================================================
# 6. Automated Strategy Backtesting Framework
# ============================================================================
build_doc(
    filename="06_automated_strategy_backtesting_framework.docx",
    title="Automated Backtesting Framework for Polymarket Trading Strategies",
    objective=(
        "Design and specify a rigorous backtesting framework that enables realistic historical "
        "simulation of Polymarket trading strategies, provides statistically valid performance "
        "metrics, and guards against overfitting through walk-forward optimization."
    ),
    prompt_text="""\
Design a comprehensive backtesting framework for the Polymarket algorithmic trading bot. The framework must support realistic replay of historical market conditions, handle the unique properties of prediction markets (binary outcomes, time-decaying optionality, event-driven resolution), and produce statistically defensible performance metrics.

1. HISTORICAL DATA COLLECTION
- Define the data pipeline for collecting and storing historical Polymarket data:
  * Gamma API: market metadata, resolution outcomes, category tags, volume, and liquidity. What endpoints to poll and at what frequency?
  * CLOB API: order book snapshots (best bid/ask, depth at each price level), trade tape (every fill with timestamp, price, size, side). How to reconstruct the full order book at any historical point in time?
  * On-chain data: CTF token transfers on Polygon for detecting large position changes not visible in the CLOB.
- Storage schema: propose a database schema (SQLite or PostgreSQL) for markets, order book snapshots, trades, and resolutions.
- Data quality: how to handle missing data, API rate limits, and market halts?

2. REPLAY ENGINE
- Design a paper-trading replay engine that simulates strategy execution against historical data:
  * The engine should process events chronologically: new market listed, order book update, trade fill, market resolution.
  * For each time step, the strategy module receives the current state (order book, portfolio, bankroll) and emits actions (place order, cancel order, do nothing).
  * Simulate realistic fills: limit orders fill only when the historical trade tape shows a trade at or through the limit price. Market orders fill at the historical best bid/ask with simulated slippage proportional to order size relative to available liquidity.
  * Account for fees (maker/taker), minimum order sizes, and the delay between order submission and fill (estimate 500ms\u20132s latency).
- How to handle strategy signals that depend on external data (e.g., NOAA forecasts)? Design a data injection interface.

3. STATISTICAL SIGNIFICANCE
- What is the minimum number of trades needed to confidently estimate the strategy's win rate? Derive using:
  * Binomial confidence intervals: for a 60% observed win rate, how many trades for the 95% CI to exclude 50%?
  * Sequential testing: is there a way to stop early if the strategy is clearly winning or losing?
- For Brier score: what is the minimum number of resolved markets for a statistically significant Brier score improvement over the market baseline?
- Account for multiple testing: if we test 10 strategy variants, how do we adjust significance thresholds (Bonferroni, FDR)?

4. WALK-FORWARD OPTIMIZATION
- Explain the walk-forward methodology: train on window W1, validate on W2, roll forward, repeat.
- For Polymarket strategies, propose specific window sizes given the available data history (approximately 12\u201318 months of liquid markets).
- What parameters should be optimized in-sample vs. fixed a priori?
  * Optimize: edge threshold for entry, Kelly fraction, category weights
  * Fix: fee assumptions, latency model, minimum bet size
- How to detect overfitting: compare in-sample Sharpe to out-of-sample Sharpe. What ratio indicates overfitting?
- Implement combinatorial purged cross-validation (CPCV) adapted for time-series prediction market data.

5. PERFORMANCE METRICS
- Define the full metrics suite for the backtester:
  * Return metrics: total return, CAGR, Sharpe ratio (adapted for prediction markets \u2014 what is the appropriate risk-free rate and return frequency?), Sortino ratio
  * Risk metrics: maximum drawdown, maximum drawdown duration, VaR (95%), CVaR (95%), probability of ruin
  * Prediction quality: Brier score over time (rolling 30-day window), calibration curve, log-loss
  * Execution quality: average slippage, fill rate, average time-to-fill
  * Strategy-specific: edge decay (how quickly does mispricing close after our entry?), win rate by category, average holding period
- Define a composite "strategy health score" that combines return, risk, and prediction quality into a single monitoring metric.

6. IMPLEMENTATION ARCHITECTURE
- Propose a Python project structure for the backtester:
  * data/ \u2014 collection scripts and database
  * engine/ \u2014 replay engine, order matching simulator
  * strategies/ \u2014 pluggable strategy modules (weather, political, sports, etc.)
  * metrics/ \u2014 performance calculation and reporting
  * reports/ \u2014 HTML/PDF report generation with charts
- How should the backtester integrate with the live trading bot? Ideally, the same strategy code runs in both backtest and live modes.
- Estimate development time for an MVP backtester vs. a production-quality framework.\
""",
    expected_outputs=[
        "Database schema diagram for historical market data storage",
        "Replay engine architecture diagram with event processing flow",
        "Minimum sample size calculations for win rate and Brier score significance",
        "Walk-forward optimization protocol with specific window sizes for Polymarket",
        "Complete performance metrics specification with formulas",
        "Overfitting detection criteria: in-sample vs. out-of-sample Sharpe ratio thresholds",
        "Python project structure with module responsibilities",
        "MVP development roadmap with estimated timeline",
        "Composite strategy health score formula and interpretation guide",
    ],
    data_sources=[
        "Polymarket Gamma API and CLOB API documentation",
        "de Prado (2018) \u2014 'Advances in Financial Machine Learning' (walk-forward, CPCV, overfitting)",
        "Bailey & de Prado (2014) \u2014 'The Deflated Sharpe Ratio' for multiple testing correction",
        "Polymarket historical trade data (CLOB trade tape)",
        "Python backtesting libraries for reference: backtrader, zipline, vectorbt",
        "SQLite / PostgreSQL documentation for schema design",
    ],
)

print("\nAll 6 research prompt documents created successfully.")
